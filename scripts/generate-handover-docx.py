"""Generate Klear Club client handover Word document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date

OUT = r"D:\klearclub\Klear Club - Site Handover Guide (Updated).docx"

doc = Document()

def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_p(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def add_number(text):
    doc.add_paragraph(text, style="List Number")

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run("Important: ")
    run.bold = True
    p.add_run(text)

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()

# ── Cover ──────────────────────────────────────────────────────────
add_title("Klear Club")
add_p("Site Handover & Self-Service Guide", bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_p(f"Prepared for Alex Delgas  |  {date.today().strftime('%B %d, %Y')}", bold=False).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
add_p("Production domain: https://klearclub.com", bold=False).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_p("Start here: Section 2 — Set up YOUR GitHub & Vercel accounts", bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ── 1. Overview ───────────────────────────────────────────────────
add_h1("1. Welcome — What You Have")
add_p(
    "Your Klear Club website is complete. This guide is written so you can manage "
    "everything yourself going forward — on your own GitHub and Vercel accounts."
)
doc.add_paragraph()
add_note(
    "The site is currently hosted on your developer's Vercel and GitHub accounts "
    "(preview: klearclub.vercel.app). Before making changes or connecting your domain, "
    "you must set up your own accounts and move the project to them. Section 2 walks "
    "you through this — it takes about 20–30 minutes."
)
doc.add_paragraph()
add_p("Once set up, you manage the site in two places:", bold=True)
add_bullet("Your GitHub account — edit prices, content, and product images in your browser")
add_bullet("Your Vercel account — hosting, deployments, and connecting klearclub.com")
doc.add_paragraph()

add_table(
    ["Area", "Status", "You manage it via"],
    [
        ["Store, product pages, cart, homepage", "Done", "Your GitHub repo"],
        ["Product images & branding", "Done", "Your GitHub repo"],
        ["Mobile design", "Done", "—"],
        ["Sale banner & promo (HEAT35)", "Done", "Your GitHub repo"],
        ["User login / OTP", "UI only", "Remedora (backend team)"],
        ["Payments", "UI only", "Remedora (backend team)"],
        ["Orders & CRM", "Placeholder", "Remedora (backend team)"],
        ["Domain klearclub.com", "Not connected yet", "Your Vercel account (Section 2 & 6)"],
    ],
)

# ── 2. Set Up YOUR Accounts (NEW — FIRST PRIORITY) ─────────────────
doc.add_page_break()
add_h1("2. Set Up YOUR GitHub & Vercel Accounts (Do This First)")
add_p(
    "The site code and hosting currently live on your developer's accounts. "
    "To own and control the site long-term, follow these steps to move everything to YOUR accounts."
)

add_h2("2.1  Create your GitHub account")
add_number("Go to https://github.com/signup")
add_number("Sign up with your business email (e.g. alex@klearclub.com or your preferred email)")
add_number("Choose a username — this becomes part of your repo URL")
add_bullet("Example: if your username is klearclub-official, your repo will be github.com/klearclub-official/klearclub")
add_number("Verify your email and complete setup")
doc.add_paragraph()

add_h2("2.2  Get the site code into YOUR GitHub")
add_p("Ask your developer to do ONE of the following (easiest first):", bold=True)
doc.add_paragraph()
add_p("Option A — Developer transfers the repo to you (recommended, easiest)", bold=True)
add_bullet("Tell your developer: \"Please transfer the klearclub repo to my GitHub account: [your username]\"")
add_bullet("Developer goes to: GitHub repo → Settings → General → Danger Zone → Transfer ownership")
add_bullet("You accept the transfer email from GitHub")
add_bullet("Repo is now yours at: github.com/YOUR-USERNAME/klearclub")
doc.add_paragraph()
add_p("Option B — Developer adds you as owner", bold=True)
add_bullet("Developer adds your GitHub account as a collaborator with Admin access")
add_bullet("You can edit everything, but repo stays under developer's account until transferred")
add_bullet("Still recommend Option A for full ownership")
doc.add_paragraph()
add_p("Option C — You import the code yourself", bold=True)
add_bullet("Ask developer to send you a ZIP of the full project folder")
add_bullet("On GitHub: click + → New repository → name it klearclub → Create")
add_bullet("Upload the ZIP files or drag-and-drop the project folders")
add_bullet("Commit — repo is now on your account")
doc.add_paragraph()

add_h2("2.3  Create your Vercel account")
add_number("Go to https://vercel.com/signup")
add_number("Sign up — choose \"Continue with GitHub\" (easiest — links accounts automatically)")
add_number("Authorize Vercel to access your GitHub account when prompted")
add_number("You now have a free Vercel account linked to your GitHub")
doc.add_paragraph()

add_h2("2.4  Deploy the site on YOUR Vercel")
add_number("Log in to https://vercel.com/dashboard")
add_number('Click "Add New…" → "Project"')
add_number("Find your klearclub repository in the list → click Import")
add_number("Vercel auto-detects Next.js — leave all settings as default:")
add_bullet("Framework Preset: Next.js")
add_bullet("Build Command: npm run build")
add_bullet("Output Directory: (leave default)")
add_bullet("Install Command: npm install")
add_number('Click "Deploy"')
add_number("Wait 2–3 minutes — Vercel builds and deploys your site")
add_number("You get a new URL like: klearclub-abc123.vercel.app (this is YOUR live preview)")
add_note("You can rename this later in Vercel → Project Settings → Domains.")
doc.add_paragraph()

add_h2("2.5  Confirm automatic deploys are working")
add_bullet("Every time you edit a file on YOUR GitHub and commit, Vercel auto-redeploys (~2 min)")
add_bullet("Test it: edit any text file on GitHub → commit → check Vercel Deployments tab → should show new deploy")
add_bullet("Bookmark your Vercel dashboard: vercel.com/dashboard")
add_bullet("Bookmark your GitHub repo: github.com/YOUR-USERNAME/klearclub")
doc.add_paragraph()

add_h2("2.6  What about the old preview URL?")
add_bullet("klearclub.vercel.app (developer's account) can stay live temporarily or be removed")
add_bullet("Once YOUR Vercel deploy is working, use your new URL for all testing")
add_bullet("Connect klearclub.com to YOUR Vercel project (Section 6) — not the developer's")
add_note("Ask your developer to remove or transfer the old Vercel project once yours is live, to avoid confusion.")

# ── 3. Your Two Dashboards ─────────────────────────────────────────
doc.add_page_break()
add_h1("3. Your Two Dashboards (Bookmark These)")
add_p("Replace YOUR-USERNAME with your actual GitHub username everywhere below.")

add_h2("GitHub — where your site code lives")
add_bullet("URL: https://github.com/YOUR-USERNAME/klearclub")
add_bullet("Use this to: change prices, edit product names, update FAQ, legal text, sale settings")
add_bullet("Every save + commit automatically triggers a new deploy on YOUR Vercel")
doc.add_paragraph()

add_h2("Vercel — where your site is hosted")
add_bullet("URL: https://vercel.com/dashboard")
add_bullet("Use this to: connect klearclub.com domain, view deployments, check build status")
add_bullet("You do NOT need Vercel for price or content changes — GitHub handles that automatically")
doc.add_paragraph()

add_h2("How updates go live (automatic)")
add_number("Edit a file on YOUR GitHub repo (in browser — no software needed)")
add_number('Click "Commit changes" at the bottom')
add_number("YOUR Vercel detects the change and rebuilds automatically (~2 minutes)")
add_number("Visit your Vercel URL (or klearclub.com once connected) and hard-refresh: Ctrl+Shift+R")

# ── 4. Pricing & Sales ─────────────────────────────────────────────
doc.add_page_break()
add_h1("4. How to Change Pricing & Sales (Step by Step)")
add_p("All pricing is in files on YOUR GitHub repo. Edit a number, commit, and the site updates automatically.")

add_h2("4.1  Change a product price")
add_number("Go to: github.com/YOUR-USERNAME/klearclub")
add_number('Navigate: src → data → products.ts → click the pencil icon (Edit)')
add_number("Find your product by slug (e.g. search for glp-3)")
add_number("Find the variant line — example:")
add_code('variant("GLP3-10", "10mg", 69.99)')
add_number("Change 69.99 to your new price")
add_number('Scroll down → "Commit changes" → confirm')
add_number("Wait ~2 minutes, hard-refresh your live site")
doc.add_paragraph()

add_h2("4.2  Change the site-wide sale percentage (currently 35% off)")
add_number("Go to: src → lib → commerce.ts → Edit")
add_code("export const SALE_OFF = 0.35;")
add_number("Change 0.35 to your new discount:")
add_bullet("25% off → 0.25  |  40% off → 0.40  |  50% off → 0.50")
add_number("Commit changes")
add_note("Also search repo for HEAT35 and 35% — update banner and product page text to match.")
doc.add_paragraph()

add_h2("4.3  Change the promo code (currently HEAT35)")
add_number("In your GitHub repo, search for: HEAT35")
add_number("Replace with your new code in these files:")
add_bullet("src/components/CartDrawer.tsx")
add_bullet("src/components/SaleBanner.tsx")
add_bullet("src/components/ProductBuyBox.tsx")
add_number("Commit each file")
doc.add_paragraph()

add_h2("4.4  Auto-apply the promo (no click required in cart)")
add_number("Go to: src → components → CartDrawer.tsx → Edit")
add_code("const [promoApplied, setPromoApplied] = useState(false);")
add_number("Change false to true → commit")
doc.add_paragraph()

add_h2("4.5  Change free shipping threshold (currently $100)")
add_number("Go to: src → lib → commerce.ts → Edit")
add_code("export const FREE_SHIPPING_AT = 100;")
add_number("Change 100 to your new threshold → commit")
doc.add_paragraph()

add_h2("4.6  Change sale banner countdown end date")
add_number("Go to: src → components → SaleBanner.tsx → Edit")
add_code('const SALE_END = new Date("2026-08-31T23:59:59");')
add_number("Change the date → commit")
doc.add_paragraph()

add_h2("4.7  Other pricing areas")
add_table(
    ["Feature", "File to edit", "Key setting"],
    [
        ["Quantity discounts (2+, 3+, 10+ bottles)", "src/lib/commerce.ts", "volumeUnitPrice()"],
        ["Build a Box (40% off)", "src/components/BuildABox.tsx", "BOX_OFF = 0.4"],
        ["Bulk Orders (40%/50% off)", "src/components/BulkOrders.tsx", "OFF_40 / OFF_50"],
        ["Research Bundles (35% off)", "src/components/ResearchBundles.tsx", "LINK_OFF = 0.35"],
        ["Membership tiers", "src/components/Membership.tsx", "TIERS array"],
    ],
)

add_h2("4.8  How the sale works today")
add_bullet("35% discount is ALREADY applied to all cart prices automatically")
add_bullet("HEAT35 in cart = confirmation step showing savings breakdown")
add_bullet("Apply does NOT change the total — it reveals retail vs. discounted price")
add_bullet("Multiple promo codes with different rules → needs Remedora backend")

# ── 5. Content Updates ─────────────────────────────────────────────
doc.add_page_break()
add_h1("5. How to Update Content Yourself")

add_table(
    ["What to change", "File in YOUR GitHub repo"],
    [
        ["Product names, descriptions, prices", "src/data/products.ts"],
        ["FAQ", "src/data/faq.ts"],
        ["COA list", "src/data/coas.ts"],
        ["Research articles", "src/data/articles.ts"],
        ["Terms of Use", "src/components/TermsPageContent.tsx"],
        ["Privacy Policy", "src/components/PrivacyPageContent.tsx"],
        ["Shipping info", "src/components/ShippingPageContent.tsx"],
        ["Returns policy", "src/components/ReturnsPageContent.tsx"],
        ["Disclaimer", "src/components/DisclaimerPageContent.tsx"],
        ["Sale banner", "src/components/SaleBanner.tsx"],
        ["Homepage featured products", "src/components/HomeFeaturedProducts.tsx"],
    ],
)

add_h2("5.1  Replace a product image")
add_number("Prepare PNG named: {product-slug}.png (e.g. glp-3.png)")
add_number("GitHub: public → products → Add file → Upload files")
add_number("Commit → hard-refresh site after ~2 min")
doc.add_paragraph()

add_h2("5.2  Add a new product")
add_number("Edit src/data/products.ts — copy existing product block, change slug/name/variants/price")
add_number("Upload image to public/products/{slug}.png")
add_number("Optionally add COA in src/data/coas.ts → commit")
doc.add_paragraph()

add_h2("5.3  Update COA PDFs")
add_number("Upload to: public/coas/pdf/{product-slug}/")
add_number("Update src/data/coas.ts → commit")

# ── 6. Domain Connection ───────────────────────────────────────────
doc.add_page_break()
add_h1("6. Connect KlearClub.com to YOUR Vercel Account")
add_note("Connect the domain to YOUR Vercel project (Section 2.4) — not the developer's old project.")

add_h2("Step 1 — Open YOUR Vercel project")
add_number("Log in to https://vercel.com/dashboard")
add_number("Open your klearclub project")
doc.add_paragraph()

add_h2("Step 2 — Add your domain")
add_number("Settings → Domains → Add")
add_number("Enter: klearclub.com → Add")
add_number("Also add: www.klearclub.com")
add_number("Vercel shows DNS records — keep this page open")
doc.add_paragraph()

add_h2("Step 3 — Configure DNS at your domain registrar")
add_p("Go to where you bought klearclub.com (GoDaddy, Namecheap, Cloudflare, etc.)")
doc.add_paragraph()
add_p("Option A — Recommended: Vercel nameservers", bold=True)
add_bullet("In registrar → Nameservers → replace with Vercel's nameservers")
add_bullet("Vercel manages everything automatically")
doc.add_paragraph()
add_p("Option B — Manual DNS records:", bold=True)
add_table(
    ["Type", "Name / Host", "Value"],
    [
        ["A", "@", "76.76.21.21"],
        ["CNAME", "www", "cname.vercel-dns.com"],
    ],
)
add_note("Confirm exact values in YOUR Vercel Domains dashboard.")
doc.add_paragraph()

add_h2("Step 4 — Verify")
add_bullet("Wait 5 min – 48 hrs (usually under 1 hour)")
add_bullet("Green checkmark in Vercel when verified")
add_bullet("HTTPS is automatic")
add_bullet("Set klearclub.com as primary domain")
add_bullet("Redirect www → klearclub.com (or vice versa)")
doc.add_paragraph()

add_h2("Step 5 — Confirm metadata on GitHub")
add_number("YOUR GitHub repo → src → app → layout.tsx → Edit")
add_code('metadataBase: new URL("https://klearclub.com")')
add_number("Should already be correct — commit only if you change it")

# ── 7. Platform & Accounts ─────────────────────────────────────────
doc.add_page_break()
add_h1("7. Platform, Login & Remedora Integration")

add_h2("7.1  Tech stack")
add_table(
    ["Item", "Detail"],
    [
        ["Framework", "Next.js 16"],
        ["Hosting", "Your Vercel account (free tier available)"],
        ["Code", "Your GitHub repo"],
        ["Backend", "Remedora team will connect"],
    ],
)

add_h2("7.2  Login & checkout — placeholder until Remedora connects")
add_table(
    ["Feature", "What users see", "Behind the scenes"],
    [
        ["/account sign-in", "Email + OTP form", "Any code accepted — no real auth"],
        ["Checkout", "Email + OTP form", "Placeholder only"],
        ["Cart", "Full cart UI", "Saved in browser only"],
        ["Google Pay", "Payment button", "No real charge"],
        ["Orders", "Order ID KC-XXXXX", "Fake — no backend"],
        ["Membership", "Pricing UI", "No billing"],
    ],
)

add_h2("7.3  Give Remedora team access")
add_p("When Remedora is ready to integrate, give them:", bold=True)
add_bullet("Access to YOUR GitHub repo (add as collaborator)")
add_bullet("The integration file: src/lib/commerce.ts")
add_bullet("Checkout payload shape (documented in that file)")
add_p("They replace the placeholder adapter with their real API for orders, auth, and payments.")

# ── 8. Deploy & Troubleshoot ─────────────────────────────────────────
add_h1("8. Deployments & Troubleshooting")
add_bullet("Commit to main branch on YOUR GitHub → YOUR Vercel auto-deploys (~2 min)")
add_bullet("Check status: YOUR Vercel dashboard → Deployments")
add_bullet("Changes not showing? Hard-refresh Ctrl+Shift+R or try incognito")
add_bullet("Deploy failed? Vercel shows the error — usually a typo in edited file")

# ── 9. Pages ───────────────────────────────────────────────────────
doc.add_page_break()
add_h1("9. All Site Pages")
add_table(
    ["Page", "URL"],
    [
        ["Homepage", "/"],
        ["Store", "/store"],
        ["Product (GLP-3 example)", "/products/glp-3"],
        ["Cart", "/cart"],
        ["Checkout", "/checkout"],
        ["Account", "/account"],
        ["Build a Box", "/build-a-box"],
        ["Bulk Orders", "/bulk"],
        ["Research Bundles", "/bundles"],
        ["Membership", "/membership"],
        ["COA Library", "/coa"],
        ["Research", "/research"],
        ["FAQ", "/faq"],
        ["Contact", "/contact"],
        ["Terms / Privacy / Shipping / Returns", "/terms, /privacy, /shipping, /returns"],
    ],
)

# ── 10. Quick Reference ──────────────────────────────────────────────
add_h1("10. Quick Reference")
add_table(
    ["I want to…", "Do this"],
    [
        ["Set up my own hosting", "Section 2 — GitHub + Vercel signup & deploy"],
        ["Change a product price", "YOUR GitHub → src/data/products.ts → commit"],
        ["Change sale %", "YOUR GitHub → src/lib/commerce.ts → SALE_OFF → commit"],
        ["Rename promo HEAT35", "YOUR GitHub → search HEAT35 → replace → commit"],
        ["Auto-show cart discount", "YOUR GitHub → CartDrawer.tsx → useState(true) → commit"],
        ["Connect klearclub.com", "YOUR Vercel → Settings → Domains → DNS"],
        ["Update FAQ", "YOUR GitHub → src/data/faq.ts → commit"],
        ["Replace product image", "YOUR GitHub → public/products/{slug}.png → upload"],
        ["Check deploy status", "YOUR Vercel → Deployments → green Ready"],
    ],
)

# ── 11. Go-Live Checklist ────────────────────────────────────────────
doc.add_page_break()
add_h1("11. Your Go-Live Checklist")
checklist = [
    ("Create YOUR GitHub account", "Section 2.1 — github.com/signup"),
    ("Get repo transferred to YOUR GitHub", "Section 2.2 — ask developer to transfer"),
    ("Create YOUR Vercel account (link to GitHub)", "Section 2.3 — vercel.com/signup"),
    ("Deploy site on YOUR Vercel", "Section 2.4 — Import repo → Deploy"),
    ("Test: edit file on GitHub → confirm Vercel redeploys", "Section 2.5"),
    ("Connect klearclub.com to YOUR Vercel", "Section 6"),
    ("Review all product prices", "Section 4.1"),
    ("Confirm sale % and promo code", "Section 4.2 & 4.3"),
    ("Give Remedora access to YOUR GitHub repo", "Section 7.3"),
    ("Ask developer to remove old Vercel preview", "Section 2.6"),
    ("Set up support@klearclub.com email", "Referenced across site"),
]
for task, where in checklist:
    p = doc.add_paragraph(style="List Number")
    p.add_run(task).bold = True
    p.add_run(f"  →  {where}")

doc.add_paragraph()
add_p("Support email on site: support@klearclub.com", bold=True)
add_p("Your GitHub: github.com/YOUR-USERNAME/klearclub")
add_p("Your Vercel: vercel.com/dashboard")
doc.add_paragraph()
add_p("— End of Guide —").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print(f"Saved: {OUT}")
